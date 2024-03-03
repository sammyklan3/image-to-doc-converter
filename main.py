import os
import glob
import pytesseract
import fitz  # PyMuPDF
import io
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

=======
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Path to the folder containing the PDF images
pdf_images_folder = "images"

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Path to the output Word document
output_docx = "output.docx"

# Function to extract text from image using OCR
def extract_text_from_image(image_path):
    try:
        return pytesseract.image_to_string(Image.open(image_path))
    except Exception as e:
        print(f"Error occurred while processing {image_path}: {str(e)}")
        return ""

# Function to extract text from PDF using PyMuPDF
def extract_text_from_pdf(pdf_path):
    texts = []
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            images = page.get_images(full=True)
            for img_index, img_info in enumerate(images):
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image = Image.open(io.BytesIO(image_bytes))
                image_path = f"temp_{page.number}_{img_index}.png"
                image.save(image_path)
                text = extract_text_from_image(image_path)
                texts.append(text)
                os.remove(image_path)  # Remove temporary image file
    except Exception as e:
        print(f"Error occurred while processing PDF: {str(e)}")
    return texts

# Function to create a Word document and format the text
def create_word_document(texts):
    doc = Document()

    # Set font style and size
    font_style = "Times New Roman"
    font_size = Pt(13)

    for text in texts:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)

        # Set font style and size
        font = run.font
        font.name = font_style
        font.size = font_size

        # Set paragraph alignment
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Set line spacing
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    try:
        doc.save("output.docx")
=======
# Function to create a Word document and format the text
def create_word_document(texts):
    doc = Document()
    for text in texts:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        font = run.font
        font.size = Pt(12)  # Change the font size as needed
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align text to left
    try:
        doc.save(output_docx)
        print("Word document created successfully.")
    except Exception as e:
        print(f"Error occurred while saving the Word document: {str(e)}")

# Main function
def main():
    source = input("Enter '1' to process a PDF file, '2' to process images directory: ")
    
    if source == '1':
        pdf_file = input("Enter the path to the PDF file: ")
        texts = extract_text_from_pdf(pdf_file)
        create_word_document(texts)
    elif source == '2':
        images_dir = input("Enter the path to the directory containing images: ")
        texts = []
        for image_file in os.listdir(images_dir):
            if image_file.endswith((".png", ".jpg", ".jpeg")):
                image_path = os.path.join(images_dir, image_file)
                text = extract_text_from_image(image_path)
                texts.append(text)
        create_word_document(texts)
    else:
        print("Invalid input. Please enter '1' or '2'.")
=======
    texts = []
    image_files = glob.glob(os.path.join(pdf_images_folder, "*.jpg")) + glob.glob(os.path.join(pdf_images_folder, "*.png"))
    if not image_files:
        print("No image files found in the specified folder.")
        return
    for image_path in image_files:
        text = extract_text_from_image(image_path)
        texts.append(text)
    create_word_document(texts)

if __name__ == "__main__":
    main()
